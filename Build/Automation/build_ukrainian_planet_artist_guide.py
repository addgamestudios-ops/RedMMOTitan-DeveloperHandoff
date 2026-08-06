from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"D:\RedMMOTitanWindowsData\ArtistHandoff\RED_Mars_27_Patches_Artist_Guide_UA.docx"
NAVY = RGBColor(20, 48, 75)
BLUE = RGBColor(35, 122, 170)
ICE = RGBColor(214, 239, 247)
ORANGE = RGBColor(196, 91, 43)
MUTED = RGBColor(90, 100, 110)
WHITE = RGBColor(255, 255, 255)


def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(115)
    p.paragraph_format.space_after = Pt(10)
    font(p.add_run(text), 29, True, NAVY)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(26)
        font(p2.add_run(subtitle), 15, False, BLUE)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    r = p.add_run(text)
    font(r, 16 if level == 1 else 13, True, NAVY if level == 1 else BLUE)
    return p


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), 11, True, NAVY)
        font(p.add_run(text[len(bold_lead):]), 11)
    else:
        font(p.add_run(text), 11)
    return p


def bullet(doc, text, numbered=False):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    font(p.add_run(text), 10.8)
    return p


def callout(doc, label, text, fill="EAF3F8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 150, 170, 150, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(label + "  "), 11, True, ORANGE)
    font(p.add_run(text), 10.8, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (header, width) in enumerate(zip(headers, widths)):
        table.columns[i].width = Inches(width)
        cell = table.rows[0].cells[i]
        shade(cell, "16304B")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        font(p.add_run(header), 9.5, True, WHITE)
    for row in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Inches(width)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            if len(table.rows) % 2 == 1:
                shade(cells[i], "F2F6F8")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(value), 9.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.45)
sec.footer_distance = Inches(0.45)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2
for style_name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 12, NAVY)):
    s = doc.styles[style_name]
    s.font.name = "Aptos"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = color

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("RED MMO  |  Planet Surface Artist Handoff"), 8.5, False, MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Марс після початку тераформування  |  UE 5.8  |  Версія для художника середовища"), 8.5, False, MUTED)

title(doc, "RED MMO: ПЛАНЕТА-МАРС", "Посібник художника середовища для 27 пласких авторських ділянок")
callout(
    doc,
    "МЕТА ПАКЕТА",
    "Надати художнику чисте, безпечне середовище для формування материків, океанів, річок, озер, пустель, полярних зон, людських і чужопланетних біомів без втручання в бойові, мережеві або UI-системи гри.",
    "E8F4F7",
)
para(doc, "Планета має окружність 50 км і є одним безперервним сферичним світом. Двадцять сім квадратних карт - це зручні пласкі джерела висоти та масок для художника. Вони не є фізичними плитами у грі: система проектує, перекриває та зшиває їх у шість безшовних граней куб-сфери.")
para(doc, "Художній задум: це Марс, який одночасно змінюють люди й чужопланетна технологія. Два колосальні тераформувальні лазери б'ють у кригу північного та південного полюсів. Тала вода формує полярні моря, океани, річкові каньйони й озера, що прорізають червону пустельну поверхню.")
doc.add_page_break()

heading(doc, "1. Що саме передано художнику", 1)
bullet(doc, "Окремий Unreal Engine 5.8 проєкт для роботи лише з поверхнею планети.")
bullet(doc, "Прототип 50-км сфери та імпортований шестигранний макрополе висоти.")
bullet(doc, "27 комплектів пласких карт висоти, суші/води, біомів і авторитету перекриття.")
bullet(doc, "Закріплена локальна версія PlanetGen, щоб результат не залежав від оновлення плагіна.")
bullet(doc, "Скрипти прийняття правок, резервного копіювання, злиття та перевірки швів.")
bullet(doc, "Цей український DOCX/PDF та короткий README з командами запуску.")
callout(doc, "НЕ ВХОДИТЬ", "Steam, PvP, UI, інвентар, персонажі, кораблі та бойовий клієнт не є частиною завдання художника. Якщо вони з'являються як залежності рівня, їх не слід редагувати.", "FFF1E8")

heading(doc, "2. Геометрія планети", 1)
data_table(doc, ["Параметр", "Значення"], [
    ("Окружність", "50 км"),
    ("Радіус базового рівня", "7,957747 км / 795 774,715 см"),
    ("Діаметр", "15,915494 км"),
    ("Площа поверхні", "795,774715 км²"),
    ("Авторські квадратні ділянки", "27"),
    ("Сховище безшовного поля", "6 квадратних граней: +X, -X, +Y, -Y, +Z, -Z"),
    ("Рівень моря", "0 см відносно планетарного датуму"),
], [2.15, 4.35])
para(doc, "Сферу неможливо без розривів скласти з 27 однакових квадратів. Тому RED розділяє авторську топологію і топологію виконання:")
bullet(doc, "Художник працює у 27 локальних пласких квадратних системах координат.", numbered=True)
bullet(doc, "Кожний квадрат проектується радіально на сферу.", numbered=True)
bullet(doc, "Сусідні квадрати перекриваються та плавно змішуються.", numbered=True)
bullet(doc, "Композиція запікається у шість точних граней куб-карти.", numbered=True)
bullet(doc, "PlanetGen ділить їх на значно менші потокові чанки для рендера і колізії.", numbered=True)
callout(doc, "ВАЖЛИВО", "Межа PNG-квадрата не є стіною, берегом, швом, зоною стримінгу або кордоном біому. У грі гравець бачить одну суцільну круглу планету.")

heading(doc, "3. Як працюють 27 пласких ділянок", 1)
para(doc, "Кожна ділянка має ID 00-26 і локальну дотичну площину. Напрямок праворуч у зображенні - локальний схід (+U). Верх зображення - локальна північ; рядки PNG зростають на південь.")
data_table(doc, ["Зона квадратної карти", "Розмір", "Роль"], [
    ("Повний квадрат", "8,000 км", "Містить ядро та перекриття з сусідами"),
    ("Повноавторитетне ядро", "5,4289168 км", "Головна зона для хабів, берегів, гір і ручної композиції"),
    ("Перо з кожного боку", "1,2855416 км", "Кубічне плавне злиття без видимого шва"),
    ("Перо при 257 x 257", "приблизно 41 піксель", "Не ставити тут різкі скелі або прямі кордони масок"),
], [1.65, 1.35, 3.5])
para(doc, "Вага перекриття обчислюється як плавність краю x кутове наближення до центру x авторитет/пріоритет. Усі активні внески нормалізуються. Це дозволяє одній ділянці м'яко контролювати важливий хаб, але не створює фізичної стіни.")

heading(doc, "4. Файли кожної ділянки", 1)
data_table(doc, ["Файл", "Формат", "Що малювати"], [
    ("RED_Patch_XX_Height_16.png", "16-bit grayscale", "Основна карта висоти"),
    ("RED_Patch_XX_Height.r16", "uint16 little-endian", "Похідний імпортний файл; зазвичай не редагувати"),
    ("RED_Patch_XX_Land.png", "8-bit grayscale", "Суша/вода: переважно 255/0, сірий лише у м'якому переході"),
    ("RED_Patch_XX_Biomes.png", "RGBA8", "R пустеля; G помірний; B холод/гори; A чужопланетний"),
    ("RED_Patch_XX_Authority.png", "8-bit grayscale", "М'який захист ручної композиції у перекритті"),
], [2.45, 1.25, 2.8])
para(doc, "Лінійне кодування висоти: 0 = -300 м, 65 535 = +300 м відносно датуму. Не переводьте 16-бітну PNG у 8-біт, палітру, JPEG, sRGB-корекцію або файл зі зміненим колірним профілем.")

heading(doc, "5. Художній напрям: тераформований Марс", 1)
para(doc, "Світ має залишатися впізнавано марсіанським: червона й охриста порода, пил, базальт, ерозійні каньйони, кратери, дюни та великі відкриті горизонти. Вода, рослинність і технологія повинні виглядати як процес, що ще триває, а не як повністю земна планета.")
bullet(doc, "Північний полюс: великий льодовий щит, зона удару першого лазера, тала вода, нові моря та холодні річкові долини.")
bullet(doc, "Південний полюс: другий лазер, розломи у кризі, парові шлейфи, льодові каньйони та інша фаза тераформування.")
bullet(doc, "Екватор і середні широти: основна марсіанська пустеля, дюни, сухі русла, солончаки та каньйони.")
bullet(doc, "Людські зони: інженерні дамби, водозбірні канали, посадкові майданчики, наукові хаби, шахти й опорні станції лазерів.")
bullet(doc, "Чужопланетні зони: неприродні кольори ґрунту, біолюмінесцентні мохи, кристалічні породи, органічні водні системи та неземні силуети рослин.")
callout(doc, "ДВА ЛАЗЕРИ", "Лазери мають читатися з поверхні й з орбіти як планетарні мегаструктури. Їхні промені спрямовані вниз у полярну кригу. Навколо точки удару потрібні світіння, пара, тріщини, водоспади талої води та великомасштабний рельєфний слід.", "E9F4FF")

heading(doc, "6. Вода, материки та дренаж", 1)
para(doc, "Вода повинна мати геологічну логіку. Спочатку визначте океанський датум і великі басейни, потім вододіли, річки, озера та дельти. Не малюйте річку як довільну синю лінію поверх рельєфу.")
bullet(doc, "Створіть два головні полярні джерела талої води.", numbered=True)
bullet(doc, "Прокладіть найнижчі шляхи через каньйони й розломи.", numbered=True)
bullet(doc, "Зберіть воду у внутрішніх озерах і кількох великих океанських басейнах.", numbered=True)
bullet(doc, "Залиште сухі старі русла поруч із активними потоками, щоб показати етапи тераформування.", numbered=True)
bullet(doc, "Перевіряйте узгодженість Height і Land на всіх перекриттях.", numbered=True)

heading(doc, "7. Безпечний цикл роботи", 1)
para(doc, "Усі команди виконуються з кореня проєкту у PowerShell.")
data_table(doc, ["Крок", "Команда / дія"], [
    ("1. Резервна точка", "python Tools\\generate_planet_authoring_patches.py snapshot-sources"),
    ("2. Редагування", "Змінити лише п'ять канонічних файлів потрібного RED_Patch_XX"),
    ("3. Прийняти", "python Tools\\generate_planet_authoring_patches.py accept-edits"),
    ("4. Перевірити детермінізм", "python Tools\\generate_planet_authoring_patches.py bake-existing"),
    ("5. Імпорт", "Запустити наданий імпортер макрополя у UE 5.8"),
    ("6. Візуальна перевірка", "Перевірити берег, колізію, шви, масштаб і вигляд з орбіти"),
], [1.75, 4.75])
callout(doc, "НЕБЕЗПЕЧНА КОМАНДА", "initialize-blockout --force-overwrite-existing-sources повністю регенерує всі 27 джерел. Не використовувати для звичайної роботи.", "FFF1E8")

heading(doc, "8. Перша практична ділянка", 1)
para(doc, "Почніть з Patch 13 біля екватора. Зробіть невелику зміну всередині ядра: одну дюнну гряду, сухе русло, базальтовий виступ або тестову зону чужопланетного ґрунту. Не торкайтеся пера під час першої вправи.")
bullet(doc, "Збережіть snapshot.")
bullet(doc, "Переконайтеся, що розмір і режим PNG не змінилися.")
bullet(doc, "Прийміть правки та повторіть bake-existing.")
bullet(doc, "Відкрийте fused prototype і порівняйте висоту, нормалі, воду та колізію.")
bullet(doc, "Лише після успішного тесту переходьте до берегової лінії між кількома ділянками.")

heading(doc, "9. Правила ручного розміщення", 1)
bullet(doc, "Великі хаби, лазерні комплекси, унікальні скелі, дамби та сюжетні POI розміщуються вручну.")
bullet(doc, "PCG використовується для заповнення, а не для заміни художньої композиції.")
bullet(doc, "Не ставте унікальний об'єкт у зоні пера, якщо він не перевірений з обох сусідніх ділянок.")
bullet(doc, "Для кожного меша перевірте масштаб у сантиметрах, pivot, collision, Nanite/LOD, матеріал і контакт із радіальною поверхнею.")
bullet(doc, "Залишайте резерв навколо хабів, доріг, берегів і посадкових зон, щоб PCG не накладав випадкові об'єкти.")

heading(doc, "10. Критерії приймання художньої ділянки", 1)
bullet(doc, "Немає квадратного контуру, сходинки або зміни масштабу на межі джерела.")
bullet(doc, "Берегова лінія збігається між Height і Land та продовжується у сусідній ділянці.")
bullet(doc, "Маски біомів змішуються природно й не утворюють рівномірної сітки.")
bullet(doc, "Гравець, наземний транспорт і корабель не зустрічають невидимої стіни чи провалу.")
bullet(doc, "Рельєф читається з землі, у польоті та з орбіти.")
bullet(doc, "Планета все ще виглядає як Марс, але вода й біоми переконливо показують людське та чужопланетне тераформування.")
bullet(doc, "Північний і південний лазерні комплекси мають різні силуети та зрозумілий зв'язок із системою води.")

heading(doc, "11. Шляхи у проєкті", 1)
data_table(doc, ["Призначення", "Шлях"], [
    ("27 джерел", "SourceArt/Planet50Km/AuthoringPatches"),
    ("Шість запечених граней", "SourceArt/Planet50Km/MacroFacesFromPatches"),
    ("Скрипт авторингу", "Tools/generate_planet_authoring_patches.py"),
    ("Профіль", "SourceArt/Planet50Km/AuthoringPatches/RED_PatchProfile.json"),
    ("Artist canvas map", "Content/RedMMO/Maps/RedPlanetGen_50km_ArtistCanvas.umap"),
    ("Fused data asset", "Content/RedMMO/Environment/DA_RED_Planet50Km_FusedHeightfield.uasset"),
], [2.05, 4.45])

heading(doc, "12. Межі цієї версії", 1)
para(doc, "Це художня робоча копія поверхні, а не збірка гри. Вона призначена для рельєфу, масок, матеріалів поверхні, біомів, водозборів і ручного розміщення. Остаточні вода, атмосфера, хмари, лазерні VFX, геймплей, реплікація та продуктивність інтегруються і приймаються у головному проєкті.")
callout(doc, "ЗОЛОТЕ ПРАВИЛО", "Спочатку snapshot, потім одна контрольована зміна, accept-edits, повторний bake, перевірка у fused map. Ніколи не просувайте художню ревізію лише тому, що PNG виглядає добре.", "E8F4F7")

doc.save(OUT)
print(OUT)
